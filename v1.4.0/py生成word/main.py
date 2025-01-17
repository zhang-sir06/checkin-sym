import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn

# 读取 CSV 文件
def read_csv(file_path):
    return pd.read_csv(file_path, encoding='utf-8-sig')

# 找出不在寝室的学生
def find_not_in_dorm(df):
    return df[df['status'] != 'on-time']

# 生成不在寝室名单 (修改)
def generate_not_in_dorm_list(df):
    teacher_data = {}
    for _, row in df.iterrows():
        teacher_name = row['teacher_name']
        class_name = row['class_name']
        member_name = row['member_name']

        if teacher_name not in teacher_data:
            teacher_data[teacher_name] = {}
        if class_name not in teacher_data[teacher_name]:
            teacher_data[teacher_name][class_name] = []
        teacher_data[teacher_name][class_name].append(member_name)
    return teacher_data

# 生成 Word 文档 (修改)
def generate_word_document(teacher_data, output_file):
    doc = Document()

    # 添加宋体字体定义到文档的styles中
    style = doc.styles.add_style('SongTi', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for teacher_name, class_data in teacher_data.items():
        doc.add_paragraph(f"{teacher_name}老师您好，女生今晚查房不在者如下：", style='SongTi')
        for class_name, member_names in class_data.items():
            member_names_str = ' '.join(member_names)
            if class_name.startswith('23'):
                color = RGBColor(255, 0, 0)  # 红色
            else:
                color = RGBColor(0, 0, 0)  # 黑色
            para = doc.add_paragraph(f"{class_name}：{member_names_str}；", style='SongTi')
            for run in para.runs:
                run.font.color.rgb = color
        doc.add_paragraph() #每个老师的信息后添加一个空行

    doc.save(output_file)

# 主函数
def main():
    csv_file_path = 'attendance_records.csv'
    output_file = '女生不在寝室名单.docx'

    # 读取数据
    df = read_csv(csv_file_path)

    # 找出不在寝室的学生
    not_in_dorm_df = find_not_in_dorm(df)

    # 生成不在寝室名单 (使用新的函数)
    not_in_dorm_list = generate_not_in_dorm_list(not_in_dorm_df)

    # 生成 Word 文档 (使用新的函数)
    generate_word_document(not_in_dorm_list, output_file)

    print(f"不在寝室名单已生成：{output_file}")

if __name__ == "__main__":
    main()